import os
import io
import shutil
import tempfile
import subprocess
from pathlib import Path

from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import RGBColor
import openpyxl
from num2words import num2words

app = Flask(__name__)

BASE_DIR = Path(__file__).parent
CONT_TEMPLATE = BASE_DIR / "CONT_plantilla.docx"
LIQUI_TEMPLATE = BASE_DIR / "LIQUI_plantilla.xlsx"


def pesos_en_letras(valor: float) -> str:
    try:
        letras = num2words(int(valor), lang="es").upper()
        return f"{letras} PESOS M/CTE"
    except Exception:
        return ""


def fmt_fecha(iso: str) -> str:
    """2026-07-04  →  4 de julio de 2026"""
    if not iso:
        return ""
    try:
        y, m, d = iso.split("-")
        meses = ["enero","febrero","marzo","abril","mayo","junio",
                 "julio","agosto","septiembre","octubre","noviembre","diciembre"]
        return f"{int(d)} de {meses[int(m)-1]} de {y}"
    except Exception:
        return iso


def set_cell(cell, text: str):
    para = cell.paragraphs[0]
    if para.runs:
        run = para.runs[0]
        for r in para.runs[1:]:
            r.text = ""
        run.text = text
    else:
        run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)


def get_unique_cells(row):
    seen, unique = set(), []
    for c in row.cells:
        if id(c) not in seen:
            seen.add(id(c))
            unique.append(c)
    return unique


def fill_contrato(doc: Document, d: dict):
    table = doc.tables[0]

    def uc(idx):
        return get_unique_cells(table.rows[idx])

    set_cell(uc(1)[0], f"NOMBRE TITULAR: {d['titular']}")
    set_cell(uc(1)[1], f"No DOCUMENTO TITULAR RSV: {d['documento']}")
    set_cell(uc(2)[0], f"NUMERO CELULAR: {d['celular']}")
    set_cell(uc(2)[1], f"CORREO ELECTRONICO: {d['correo']}")
    set_cell(uc(3)[0], f"HOTEL: {d['hotel']}")
    set_cell(uc(3)[1], f"ORIGEN: {d['origen']}")
    set_cell(uc(3)[2], f"DESTINO: {d['destino']}")
    set_cell(uc(4)[0], f"ADULTOS: {d['adultos']}")
    set_cell(uc(4)[1], f"FECHA DE SALIDA: {fmt_fecha(d['fecha_ida'])}")
    set_cell(uc(5)[0], f"NIÑOS 2-11 AÑOS: {d['ninos']}")
    set_cell(uc(5)[1], f"FECHA DE REGRESO: {fmt_fecha(d['fecha_regreso'])}")
    set_cell(uc(6)[0], f"INFANTES 0-23 MESES: {d['infantes']}")
    set_cell(uc(6)[1], f"TOTAL NOCHES: {d['noches']}")
    set_cell(uc(7)[0], f"ACOMODACION HABITACIONAL: {d['acomodacion']}")

    for i, p in enumerate(d.get("pasajeros", [])[:7]):
        cells = get_unique_cells(table.rows[10 + i])
        set_cell(cells[0], p.get("nombre", ""))
        set_cell(cells[1], p.get("doc", ""))

    tarifa = float(d.get("valor_total", 0) or 0)
    cuota  = float(d.get("cuota_inicial", 0) or 0)
    set_cell(uc(17)[0], f"TARIFA TOTAL EN LETRAS: {pesos_en_letras(tarifa)}")
    set_cell(uc(17)[1], f"TARIFA EN NUMEROS: ${tarifa:,.0f}")
    set_cell(uc(18)[0], f"TARIFA: ${tarifa:,.0f}")
    if len(uc(19)) >= 3:
        set_cell(uc(19)[2], f"CUOTA INICIAL: ${cuota:,.0f}")
    if len(uc(20)) >= 3:
        set_cell(uc(20)[2], f"NUMERO CUOTAS: {d.get('num_cuotas','')}")
    if len(uc(21)) >= 3:
        set_cell(uc(21)[2], f"VALOR CUOTA: ${float(d.get('valor_cuota',0) or 0):,.0f}")


def run_soffice(args):
    env = os.environ.copy()
    env["SAL_USE_VCLPLUGIN"] = "svp"
    subprocess.run(["soffice"] + args, env=env, check=True,
                   capture_output=True, timeout=90)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/contrato", methods=["POST"])
def generar_contrato():
    d = request.json or {}
    output_format = d.get("formato", "pdf")
    nombre_safe = (d.get("titular") or "cliente").replace(" ", "_")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        docx_out = tmp / f"Contrato_{nombre_safe}.docx"
        shutil.copy(CONT_TEMPLATE, docx_out)

        doc = Document(docx_out)
        fill_contrato(doc, d)
        doc.save(docx_out)

        if output_format == "docx":
            buf = io.BytesIO(docx_out.read_bytes())
            buf.seek(0)
            return send_file(buf,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True, download_name=docx_out.name)

        run_soffice(["--headless", "--convert-to", "pdf",
                     "--outdir", str(tmp), str(docx_out)])
        pdf_out = docx_out.with_suffix(".pdf")
        buf = io.BytesIO(pdf_out.read_bytes())
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True,
                         download_name=f"Contrato_{nombre_safe}.pdf")


@app.route("/liquidacion", methods=["POST"])
def generar_liquidacion():
    d = request.json or {}
    output_format = d.get("formato", "pdf")
    nombre_safe = (d.get("titular") or "cliente").replace(" ", "_")

    tarifa_sencilla = float(d.get("tarifa_sencilla", 0) or 0)
    tarifa_doble    = float(d.get("tarifa_doble", 0) or 0)
    tarifa_triple   = float(d.get("tarifa_triple", 0) or 0)
    tarifa_nino     = float(d.get("tarifa_nino", 0) or 0)
    tarifa_infante  = float(d.get("tarifa_infante", 0) or 0)
    cant_sencilla   = int(d.get("cant_sencilla", 0) or 0)
    cant_doble      = int(d.get("cant_doble", 0) or 0)
    cant_triple     = int(d.get("cant_triple", 0) or 0)
    cant_nino       = int(d.get("cant_nino", 0) or 0)
    cant_infante    = int(d.get("cant_infante", 0) or 0)
    tours           = [float(x) for x in d.get("tours", [0]*5)]
    asistencia      = [float(x) for x in d.get("asistencia", [0]*5)]
    equipaje        = [float(x) for x in d.get("equipaje", [0]*5)]
    cuota_inicial   = float(d.get("cuota_inicial", 0) or 0)
    num_cuotas      = int(d.get("num_cuotas", 1) or 1)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        xlsx_out = tmp / f"Liquidacion_{nombre_safe}.xlsx"
        shutil.copy(LIQUI_TEMPLATE, xlsx_out)

        wb = openpyxl.load_workbook(xlsx_out)
        ws = wb.active

        cell_map = {
            "D5":  "Yamile Segura",
            "D6":  d.get("asesor", ""),
            "D7":  fmt_fecha(d.get("fecha_rsv", "")),
            "D9":  d.get("titular", ""),
            "D10": d.get("documento", ""),
            "H9":  d.get("celular", ""),
            "H11": d.get("correo", ""),
            "D13": d.get("origen", ""),
            "H13": d.get("destino", ""),
            "D14": fmt_fecha(d.get("fecha_ida", "")),
            "H14": fmt_fecha(d.get("fecha_regreso", "")),
            "D15": int(d.get("adultos", 0) or 0),
            "F15": int(d.get("ninos", 0) or 0),
            "H15": int(d.get("infantes", 0) or 0),
            "D16": d.get("hotel", ""),
            "F17": d.get("noches", 0),
            "H17": d.get("tipo_vuelo", "Comercial"),
            "D19": cant_sencilla, "E19": cant_doble,
            "F19": cant_triple,   "G19": cant_nino, "H19": cant_infante,
            "D21": tarifa_sencilla, "E21": tarifa_doble,
            "F21": tarifa_triple,   "G21": tarifa_nino, "H21": tarifa_infante,
            "D22": tours[0], "E22": tours[1], "F22": tours[2],
            "G22": tours[3], "H22": tours[4],
            "D23": asistencia[0], "E23": asistencia[1], "F23": asistencia[2],
            "G23": asistencia[3], "H23": asistencia[4],
            "D24": equipaje[0], "E24": equipaje[1],
            "F24": equipaje[2],  "H24": equipaje[4] if len(equipaje) > 4 else 0,
            "D25": cant_sencilla, "E25": cant_doble,
            "F25": cant_triple,   "G25": cant_nino, "H25": cant_infante,
        }
        for ref, val in cell_map.items():
            ws[ref] = val

        total = (tarifa_sencilla*cant_sencilla + tarifa_doble*cant_doble +
                 tarifa_triple*cant_triple + tarifa_nino*cant_nino +
                 tarifa_infante*cant_infante)
        # Si no hay tarifas por acomodación, usar valor_total del form
        if total == 0:
            total = float(d.get("valor_total", 0) or 0)

        ws["D28"] = pesos_en_letras(total)
        ws["D30"] = cuota_inicial
        ws["F30"] = num_cuotas
        ws["H30"] = total - cuota_inicial

        wb.save(xlsx_out)

        if output_format == "xlsx":
            buf = io.BytesIO(xlsx_out.read_bytes())
            buf.seek(0)
            return send_file(buf,
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                as_attachment=True, download_name=xlsx_out.name)

        run_soffice(["--headless", "--convert-to", "pdf",
                     "--outdir", str(tmp), str(xlsx_out)])
        pdf_out = xlsx_out.with_suffix(".pdf")
        buf = io.BytesIO(pdf_out.read_bytes())
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True,
                         download_name=f"Liquidacion_{nombre_safe}.pdf")


if __name__ == "__main__":
    app.run(debug=True, port=5000)
